import os
import re
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document

def parse_file_content(file_path, file_name):
    """
    Faylni matn ko'rinishiga o'tkazish
    PDF, DOCX, CSV, TXT fayllarini qo'llab-quvvatlaydi
    """
    file_extension = os.path.splitext(file_name)[1].lower()
    
    if file_extension == '.pdf':
        return parse_pdf(file_path)
    elif file_extension == '.docx':
        return parse_docx(file_path)
    elif file_extension == '.csv':
        return parse_csv(file_path)
    elif file_extension == '.txt':
        return parse_txt(file_path)
    else:
        raise Exception("Qo'llab-quvvatlanmaydigan fayl turi")

def parse_pdf(file_path):
    """PDF faylini matnga aylantiradigan funksiya"""
    try:
        reader = PdfReader(file_path)
        text = ""
        
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            page_text = page.extract_text()
            text += (page_text or "")
        
        # Matnni tozalash
        text = clean_text(text)
        if not text.strip():
            raise Exception("PDF faylidan matn ajratib olib bo'lmadi (ehtimol rasmli fayl)")
        return text
    
    except Exception as e:
        raise Exception(f"PDF faylini o'qishda xatolik: {str(e)}")

def parse_docx(file_path):
    """DOCX faylini matnga aylantiradigan funksiya"""
    try:
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Jadvallarni ham o'qish
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        # Matnni tozalash
        text = clean_text(text)
        return text
    
    except Exception as e:
        raise Exception(f"DOCX faylini o'qishda xatolik: {str(e)}")

def parse_csv(file_path):
    """CSV faylini matnga aylantiradigan funksiya"""
    try:
        # CSV faylini o'qish
        df = pd.read_csv(file_path)
        
        # Ma'lumotlarni matn ko'rinishiga o'tkazish
        text = "CSV fayli ma'lumotlari:\n\n"
        
        # Ustun nomlarini qo'shish
        text += "Ustunlar: " + ", ".join(df.columns.tolist()) + "\n\n"
        
        # Ma'lumotlarni qo'shish
        for index, row in df.iterrows():
            text += f"Qator {int(index) + 1}:\n"
            for col in df.columns:
                text += f"  {col}: {row[col]}\n"
            text += "\n"
        
        return text
    
    except Exception as e:
        raise Exception(f"CSV faylini o'qishda xatolik: {str(e)}")

def parse_txt(file_path):
    """TXT faylini o'qiydigan funksiya"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        # Matnni tozalash
        text = clean_text(text)
        return text
    
    except UnicodeDecodeError:
        # UTF-8 bilan o'qib bo'lmasa, boshqa encoding bilan sinab ko'rish
        try:
            with open(file_path, 'r', encoding='cp1251') as file:
                text = file.read()
            return clean_text(text)
        except:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                return clean_text(text)
            except Exception as e:
                raise Exception(f"TXT faylini o'qishda xatolik: {str(e)}")
    
    except Exception as e:
        raise Exception(f"TXT faylini o'qishda xatolik: {str(e)}")

def clean_text(text):
    """Matnni tozalash va formatlash"""
    if not text:
        return ""
    
    # Ortiqcha bo'shliqlarni olib tashlash
    text = re.sub(r'\s+', ' ', text)
    
    # Bosh va oxiridagi bo'shliqlarni olib tashlash
    text = text.strip()
    
    return text

def is_allowed_file(filename):
    """Ruxsat etilgan fayl turlarini tekshirish"""
    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.csv', '.txt'}
    return os.path.splitext(filename)[1].lower() in ALLOWED_EXTENSIONS

def get_file_size_mb(file_path):
    """Fayl hajmini MB da qaytarish"""
    size_bytes = os.path.getsize(file_path)
    size_mb = size_bytes / (1024 * 1024)
    return round(size_mb, 2)